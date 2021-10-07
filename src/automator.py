import os
import locale
from datetime import datetime
import json

import PyPDF2
import docx
import docx2pdf
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

#from .elem import Person, Organization
import elem
import documents as docs

# References
# PDF Editing : https://caendkoelsch.wordpress.com/2019/05/10/merging-multiple-pdfs-into-a-single-pdf/
# Word Editing: https://tech-cookbook.com/2019/10/21/how-to-work-with-docx-in-python/

STD_DIR = os.path.join(os.getcwd(), "src",
                       "std_application")
STD_DIR_CURRICULUM_VITAE = os.path.join(STD_DIR, "1_curriculum_vitae")
STD_DIR_CERTIFICATES = os.path.join(STD_DIR, "2_certificates")
STD_FILE_CERTICATES_INDICES = os.path.join(STD_DIR_CERTIFICATES, "index.json")
STD_APPLICATION_PHOTO = os.path.join(STD_DIR, "max_mustermann.png")


MAIN_DIR = os.getcwd()





class Field:

    def __init__(self,
                 identifier,
                 style):
        self.identifier = identifier
        self.style = style


class Job:

    def __init__(self,
                 position,
                 identification_number,
                 link=None):
        self.position = position
        self.identification_number = identification_number
        self.link = link

    def print(self):
        pass

class JobApplicationFactory:
    def __init__(self):
        self._builders = {}

    def register_builder(self, key, builder):
        self._builders[key] = builder

    def create(self, key, **kwargs):
        builder = self._builders.get(key)
        if not builder:
            raise ValueError(key)
        return builder(**kwargs)

class JobApplication:

    def __init__(self,
                 candidate,
                 their_contact_person,
                 their_organization,
                 their_job_offer):

        self.fields = {}
        self.fields_img = {}

        #self.application_language = application_language
        self.fields["numbering"] = ""
        self.fields["candidate.name"] = candidate.name
        self.fields["candidate.address"] = candidate.address
        self.fields["candidate.postcode"] = candidate.postcode
        self.fields["candidate.city"] = candidate.city
        self.fields["candidate.telephone_number"] = candidate.telephone_number
        self.fields["candidate.email_address"] = candidate.email_address

        self.fields["their_contact_person.name"] = their_contact_person.name
        self.fields["contact_surname"] = their_contact_person.surname
        self.contact_gender = their_contact_person.gender

        # Organization
        self.fields["their_organization.name"] = their_organization.name
        self.fields["organization_name_abbreviation"] = their_organization.name
        self.fields["their_organization.address"] = their_organization.address
        self.fields["their_organization.postcode"] = their_organization.postcode
        self.fields["their_organization.city"] = their_organization.city

        # Job offer
        self.fields["job_position"] = their_job_offer.position
        self.fields["job_identification_number"] = their_job_offer.identification_number

        self.dirname = "{}_{}_{}".format(self.fields["job_identification_number"],
                                         self.fields["organization_name_abbreviation"],
                                         self.fields["job_position"].replace(' ', '-'))

        self.dirpath_application = {"main": os.path.join(MAIN_DIR, self.dirname),
                                    "motivation_letter": os.path.join(MAIN_DIR, self.dirname, "0_motivation_letter"),
                                    "curriculum_vitae": os.path.join(MAIN_DIR, self.dirname, "1_curriculum_vitae"),
                                    "certificates": os.path.join(MAIN_DIR, self.dirname, "2_certificates"),
                                    "corrections": os.path.join(MAIN_DIR, self.dirname, "3_corrections"),
                                    "assembly": os.path.join(MAIN_DIR, self.dirname, "4_assembly")}

        self.filepaths_application = {"curriculum_vitae_docx": os.path.join(self.dirpath_application["curriculum_vitae"], "curriculum_vitae.docx"),
                                      "curriculum_vitae_pdf": os.path.join(self.dirpath_application["curriculum_vitae"], "curriculum_vitae.pdf"),
                                      "motivation_letter_docx": os.path.join(self.dirpath_application["motivation_letter"], "motivation_letter.docx"),
                                      "motivation_letter_pdf": os.path.join(self.dirpath_application["motivation_letter"], "motivation_letter.pdf"),
                                      "intro_docx": os.path.join(self.dirpath_application["certificates"], "intro.docx"),
                                      "intro_pdf": os.path.join(self.dirpath_application["certificates"], "intro.pdf")}

        self.fields_img["picture"] = STD_APPLICATION_PHOTO

        if os.path.isdir(self.dirpath_application["main"]) is False:
            for key, path in self.dirpath_application.items():
                os.makedirs(path)

    def export(self):

        def find_replace(paragraph_keyword, draft_keyword, paragraph, style):
            if paragraph_keyword in paragraph.text:
                # print("found")
                if draft_keyword is None:
                    draft_keyword = ""
                try:
                    paragraph.text = paragraph.text.replace('#{}#'.format(paragraph_keyword), draft_keyword)
                    paragraph.style = style
                except:
                    print("Keyword: {}, Value: {}".format(paragraph_keyword,draft_keyword))
                    ValueError()

        def find_replace_picture(paragraph_keyword, draft_keyword, paragraph):
            if paragraph_keyword in paragraph.text:
                # print("found")
                #try:
                    #paragraph.text = paragraph.text.replace('#{}#'.format(paragraph_keyword), "")
                paragraph.clear()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = paragraph.add_run()
                inline_shape = run.add_picture(draft_keyword)
                #inline_shape = grun.add_picture(draft_keyword, 0.6*3556000, 0.75*3556000)
                #inline_shape.width =
                #print("Hello")
                #inline_shape.width = 10
                #inline_shape.height = 10
                #except:
                #    print("Keyword: {}, Value: {}".format(paragraph_keyword,draft_keyword))
                #    ValueError()

        #curriculum_vitae_docx.paragraphs[0].runs[0].add_picture(STD_APPLICATION_PHOTO)
        #curriculum_vitae_docx.paragraphs[1].runs[0].add_picture(STD_APPLICATION_PHOTO)



        if hasattr(self, "std_intro") is False:
            AttributeError("No reference document found.")

        ## CURRICULUM VITAE
        #curriculum_vitae_docx = Document(self.std["curriculum_vitae"])
        #curriculum_vitae_docx.save(self.filepaths_application["curriculum_vitae_docx"])
        #docx2pdf.convert(self.filepaths_application["curriculum_vitae_docx"],
        #                 self.filepaths_application["curriculum_vitae_pdf"])


        # MOTIVATION LETTER
        motivation_letter_docx = docx.Document(self.std_intro)

        style = motivation_letter_docx.styles.add_style('Original', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        font.name = 'Arial'
        font.size = docx.shared.Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.15
        paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.MULTIPLE

        style_w = motivation_letter_docx.styles.add_style('White', WD_STYLE_TYPE.PARAGRAPH)
        font = style_w.font
        font.name = 'Arial'
        font.size = docx.shared.Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.15
        paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.MULTIPLE
        font.color.rgb = docx.shared.RGBColor(255, 255, 255)
        font.bold = True

        stylel = motivation_letter_docx.styles["List Number"]
        font = stylel.font
        font.name = 'Arial'
        font.size = docx.shared.Pt(11)
        paragraph_format = stylel.paragraph_format
        tb = paragraph_format.tab_stops
        paragraph_format.line_spacing = 1.15
        paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.MULTIPLE
        font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        #header = motivation_letter_docx.sections[0].header
        header = motivation_letter_docx.sections[0].first_page_header
        for paragraph in header.paragraphs:
            for key, item in self.fields.items():
                if key == "job_position" or key == "job_identification":
                    find_replace(key, item, paragraph, style_w)
                else:
                    find_replace(key, item, paragraph, style)

        for paragraph in motivation_letter_docx.paragraphs:
            for key, item in self.fields.items():
                if key == "job_position" or key == "job_identification":
                    find_replace(key, item, paragraph, style_w)
                else:
                    find_replace(key, item, paragraph, style)

        motivation_letter_docx.save(self.filepaths_application["motivation_letter_docx"])

        # Curriculum Vitae
        curriculum_vitae_docx = docx.Document(self.std["curriculum_vitae"])
        for paragraph in curriculum_vitae_docx.paragraphs:
            for key, item in self.fields_img.items():
                find_replace_picture(key, item, paragraph)

        for paragraph in curriculum_vitae_docx.paragraphs:
            for key, item in self.fields.items():
                if key == "job_position" or key == "job_identification":
                    find_replace(key, item, paragraph, style_w)
                else:
                    find_replace(key, item, paragraph, style)

        curriculum_vitae_docx.save(self.filepaths_application["curriculum_vitae_docx"])

        # PART 2: Convert DOC file to PDF file
        docx2pdf.convert(self.filepaths_application["curriculum_vitae_docx"],
                         self.filepaths_application["curriculum_vitae_pdf"])


        # CERTIFICATES
        document = docx.Document(self.std_intro)

        style = document.styles.add_style('Original', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        font.name = 'Arial'
        font.size = docx.shared.Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.15
        paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.MULTIPLE

        style_w = document.styles.add_style('White', WD_STYLE_TYPE.PARAGRAPH)
        font = style_w.font
        font.name = 'Arial'
        font.size = docx.shared.Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.15
        paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.MULTIPLE
        font.color.rgb = docx.shared.RGBColor(255, 255, 255)
        font.bold = True

        stylel = document.styles["List Number"]
        font = stylel.font
        font.name = 'Arial'
        font.size = docx.shared.Pt(11)
        paragraph_format = stylel.paragraph_format
        tb = paragraph_format.tab_stops
        paragraph_format.line_spacing = 1.15
        paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.MULTIPLE
        font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        header = document.sections[0].header
        for paragraph in header.paragraphs:
            for key, item in self.fields.items():
                if key == "job_position" or key == "job_identification":
                    find_replace(key, item, paragraph, style_w)
                else:
                    find_replace(key, item, paragraph, style)

        for paragraph in document.paragraphs:
            for key, item in self.fields.items():
                if key == "job_position" or key == "job_identification":
                    find_replace(key, item, paragraph, style_w)
                elif key == "numbering" and key in paragraph.text:
                    total_number_of_pages = 2
                    for certificate in self.certificates:
                        paragraph.insert_paragraph_before('{}\x09{}'.format(certificate["name"],
                                                                            total_number_of_pages),
                                                          style='List Number')
                        total_number_of_pages += certificate["number_of_pages"]
                    find_replace(key, item, paragraph, style)
                else:
                    find_replace(key, item, paragraph, style)

        document.save(self.filepaths_application["intro_docx"])

        # PART 2: Convert DOC file to PDF file
        docx2pdf.convert(self.filepaths_application["intro_docx"], self.filepaths_application["intro_pdf"])

        # PART 3. Merge certificates together
        pdf_certificates = PyPDF2.PdfFileMerger()
        pdf_certificates.append(self.filepaths_application["intro_pdf"])
        for certificate in self.certificates:
            pdf_certificates.append(certificate["filepath"])
        pdf_certificates.write(os.path.join(self.dirpath_application["certificates"], "certificates.pdf"))
        pdf_certificates.close()

class JobApplicationEnglish(JobApplication):

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

    def set_fields(self):
        locale.setlocale(locale.LC_TIME, 'en_US')
        self.std_intro = os.path.join(os.getcwd(), STD_DIR, "english", "intro_en.docx")
        self.fields["date"] = datetime.now().strftime("%d. %B %Y")
        # Check if there is a job identification number
        if self.fields["job_identification_number"] is None:
            self.fields["job_identification"] = "from your website"
        else:
            self.fields["job_identification"] = "with code number {}".format(self.fields["job_identification_number"])
        if self.contact_gender == "Male":
            self.fields["form_of_address"] = "Mr."
            self.fields["greeting"] = "Dear"
        elif self.contact_gender == "Female":
            self.fields["form_of_address"] = "Ms."
            self.fields["greeting"] = "Diverse"
        elif self.contact_gender is None:
            self.fields["greeting"] = "Dear"
            self.fields["form_of_address"] = "all"

class JobApplicationGerman(JobApplication):

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

    def set_fields(self):
        locale.setlocale(locale.LC_TIME, 'de_DE')
        self.fields["date"] = datetime.now().strftime("%d. %B %Y")
        self.std_intro = os.path.join(os.getcwd(), STD_DIR, "2_certificates", "intro_de.docx")
        self.std = {}
        self.std["motivation_letter"] = os.path.join(os.getcwd(), STD_DIR, "0_motivation_letter", "motivation_letter_de.docx")
        self.std["curriculum_vitae"] = os.path.join(os.getcwd(), STD_DIR, "1_curriculum_vitae", "curriculum_vitae_de.docx")

        if self.fields["job_identification_number"] is None:
            self.fields["job_identification"] = "Ihrer Website"
        else:
            self.fields["job_identification"] = "mit der Kennziffer {}".format(self.fields["job_identification_number"])

        # Check contact person gender
        if self.contact_gender == "Male":
            self.fields["form_of_address"] = "Herr"
            self.fields["greeting"] = "Sehr geehrter"
            self.fields["filler_greeting"] = " "
        elif self.contact_gender == "Female":
            self.fields["form_of_address"] = "Frau"
            self.fields["greeting"] = "Sehr geehrte"
            self.fields["filler_greeting"] = " "
        elif self.contact_gender is None:
            self.fields["form_of_address"] = ""
            self.fields["contact_surname"] = "Damen und Herren"
            self.fields["greeting"] = "Sehr geehrte"
            self.fields["filler_greeting"] = ""

        self.database_certificates = []
        for file in os.listdir(STD_DIR_CERTIFICATES):
                if file.endswith(".pdf"):
                    this_certificate = docs.Certificate(file)
                    this_certificate.get_name(language="German")
                    self.database_certificates.append({"name": this_certificate.name,
                                                       "number_of_pages": this_certificate.number_of_pages,
                                                       "filename": this_certificate.filename,
                                                       "filepath": this_certificate.filepath})

        f = open(STD_FILE_CERTICATES_INDICES,)
        indices = json.load(f)
        self.default_list = [values for key, values in indices.items() if key == "German"][0]["default_list"]

        self.certificates = []
        for filename in self.default_list:
            self.certificates.append([certificate for certificate in self.database_certificates if certificate["filename"] == filename][0])

if __name__ == '__main__':

    factory = JobApplicationFactory()
    factory.register_builder("English", JobApplicationEnglish)
    factory.register_builder("German", JobApplicationGerman)

    me = elem.Person(name="Max Mustermann",
                gender="Male",
                address="Musterstraße 12",
                postcode="12345",
                city="Musterstadt",
                telephone_number="01234 5789",
                email_address="max.mustermann@bewerbung.co")

    their_contact_person = elem.Person(name="Musterfrau",
                                  gender="Female")

    their_organization = elem.Organization(name="Musterfirma GmbH",
                                      name_abbreviation="Musterfima GmbH",
                                      address="Musterstraße 11",
                                      postcode="12345",
                                      city="Musterstadt")

    their_job_offer = Job(position="Musterberuf",
                          identification_number="9999-9999")

    app_config = {"candidate": me,
                  "their_contact_person": their_contact_person,
                  "their_organization": their_organization,
                  "their_job_offer": their_job_offer}

    my_job_application = factory.create("German", **app_config)
    my_job_application.set_fields()
    my_job_application.export()

